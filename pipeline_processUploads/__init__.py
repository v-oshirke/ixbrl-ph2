import azure.functions as func
import logging
from docx import Document
import fitz
from utils.blob_functions import list_blobs, get_blob_content, write_to_blob
from utils import get_month_date
import io
import os
import json
import base64

# Libraries used in the future Document Processing client code
from azure.identity import DefaultAzureCredential
from azure.ai.documentintelligence import DocumentIntelligenceClient
from azure.ai.documentintelligence.models import AnalyzeResult, AnalyzeDocumentRequest

# Variables used by Document Processing client code
endpoint =os.getenv("AIMULTISERVICES_ENDPOINT") # Add the AI Services Endpoint value from Azure Function App settings

def extract_text_from_blob(blob_name):
    try:
        credential = DefaultAzureCredential()                
        client = DocumentIntelligenceClient(
            endpoint=endpoint, credential=credential
        )
        
        content = get_blob_content("bronze", blob_name)
        
        base64_content = base64.b64encode(content).decode('utf-8')    

        poller = client.begin_analyze_document(
            # AnalyzeDocumentRequest Class: https://learn.microsoft.com/en-us/python/api/azure-ai-documentintelligence/azure.ai.documentintelligence.models.analyzedocumentrequest?view=azure-python
            "prebuilt-read", AnalyzeDocumentRequest(bytes_source=base64_content
        ))
        result: AnalyzeResult = poller.result()
        
        if result.paragraphs:    
            paragraphs = "\n".join([paragraph.content for paragraph in result.paragraphs])            
        
        return paragraphs
        
    except Exception as e:
        logging.error(f"Error processing {blob_name}: {e}")
        return None

# def extract_text_from_docx(blob_name):
#     try:
#         # Get the content of the blob
#         content = get_blob_content("bronze", blob_name)
#         # Load the content into a Document object
#         doc = Document(io.BytesIO(content))
#         # Extract and print the text
#         full_text = []
#         for paragraph in doc.paragraphs:
#             full_text.append(paragraph.text)

#         # Combine paragraphs into a single string
#         text = "\n".join(full_text)
#         return text
#     except Exception as e:
#         logging.error(f"Error processing {blob_name}: {e}")
#         return None

# def extract_text_from_pdf(blob_name):
#     try:
#         # Get the content of the blob
#         content = get_blob_content("bronze", blob_name)
#         # Load the PDF document
#         doc = fitz.open(stream=content, filetype="pdf")
#         # Extract text from all pages
#         text = "\n".join(page.get_text() for page in doc)
#         return text
#     except Exception as e:
#         logging.error(f"Error processing {blob_name}: {e}")
#         return None

def main(req: func.HttpRequest) -> func.HttpResponse:
    logging.info('Python HTTP trigger function processed a request.')
    
    req_body = req.get_json()
    # Get the list of blobs sent from the frontend (if any)
    selected_blobs = req_body.get("blobs", None)
    
    processed_files = []
    errors = []
    
    # Document Intelligence supported suffixes
    suffixes = (".jpg", ".jpeg", ".png", ".tiff", ".docx", ".xlsx", ".pptx", ".pdf")

    # Lists blobs in the 'bronze' container
    if selected_blobs:
        for blob in selected_blobs:
            if blob.get("container") != "bronze":
                logging.info(f"Skipping blob not in bronze container: {blob}")
                continue
            try:
                blob_name = blob.get("name")

                # Extract text from supported file types using Document Intelligence
                if blob_name.endswith(suffixes):
                    logging.info(f"Processing: {blob_name}")
                    text = extract_text_from_blob(blob_name)
                    if text:
                        sourcefile = os.path.splitext(os.path.basename(blob_name))[0]
                        write_to_blob(f"silver", f"{sourcefile}.txt", text)
                        processed_files.append(blob_name)
                    else:
                        errors.append(f"Failed to extract text from: {blob_name}")

                # if blob_name.endswith(".docx"):
                #     logging.info(f"Processing DOCX: {blob_name}")
                #     text = extract_text_from_docx(blob_name)
                #     if text:
                #         sourcefile = os.path.splitext(os.path.basename(blob_name))[0]
                #         write_to_blob(f"silver", f"{sourcefile}.txt", text)
                #         processed_files.append(blob_name)
                #     else:
                #         errors.append(f"Failed to extract text from DOCX: {blob_name}")

                # elif blob_name.endswith(".pdf"):
                #     logging.info(f"Processing PDF: {blob_name}")
                #     text = extract_text_from_pdf(blob_name)
                #     if text:
                #         sourcefile = os.path.splitext(os.path.basename(blob_name))[0]
                #         write_to_blob(f"silver", f"{sourcefile}.txt", text)
                #         processed_files.append(blob_name)
                #     else:
                #         errors.append(f"Failed to extract text from PDF: {blob_name}")

                else:
                    logging.info(f"Skipping unsupported file type: {blob_name}")
                    errors.append(f"Unsupported file type: {blob_name}")

            except Exception as e:
                error_message = f"Error processing {blob_name}: {e}"
                logging.error(error_message)
                errors.append(error_message)
    else:
        return func.HttpResponse(
            json.dumps({"error": "No blobs provided."}),
            status_code=400,
            mimetype="application/json"
        )
    # Prepare the response payload
    response_data = {
        "processedFiles": processed_files,
        "errors": errors,
        "status": "started" if not errors else "completed_with_errors"
    }

    # Return appropriate status code based on the outcome
    if errors:
        return func.HttpResponse(
            json.dumps(response_data),
            status_code=500,  # Internal Server Error
            mimetype="application/json"
        )
    else:
        return func.HttpResponse(
            json.dumps(response_data),
            status_code=200,  # Success
            mimetype="application/json"
        )